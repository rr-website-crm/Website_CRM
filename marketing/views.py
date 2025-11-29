from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.utils import timezone
from django.http import JsonResponse
from django.views.decorators.http import require_http_methods
from django.core.exceptions import ValidationError
from django.db import transaction
from django.core.paginator import Paginator
from django.urls import reverse
from contextlib import contextmanager
from django.conf import settings
import json
import time
import os
from openai import OpenAI
try:
    from bson.decimal128 import Decimal128
except ImportError:
    Decimal128 = None
from .models import Job, JobAttachment, JobSummaryVersion, JobActionLog, log_job_activity
from .models import Customer, CustomerActionLog
from accounts.models import ActivityLog, CustomUser
from accounts.services import log_activity_event
import logging

from superadminpanel.models import (
    TemplateMaster, ProjectGroupMaster, PriceMaster,
    ReferencingMaster, AcademicWritingMaster,
    JobTemplate, JobTask
)
from datetime import timedelta

logger = logging.getLogger('marketing')


def _decimal_to_float(value):
    """Convert Decimal128/Decimal/string to float safely."""
    if value is None:
        return None
    if Decimal128 and isinstance(value, Decimal128):
        try:
            return float(value.to_decimal())
        except Exception:
            return None
    try:
        return float(value)
    except (TypeError, ValueError):
        return None

def _to_float(value, default=0.0):
    """Convert values to float, tolerating Decimal128 and returning a default on failure."""
    if value is None:
        return default
    if Decimal128 and isinstance(value, Decimal128):
        try:
            return float(value.to_decimal())
        except Exception:
            return default
    try:
        return float(value)
    except (TypeError, ValueError):
        return default

def _to_price_master_level(level_value):
    """Map job-friendly level values to PriceMaster stored values."""
    if not level_value:
        return None
    normalized = str(level_value).strip().lower()
    if normalized == 'advanced':
        return 'advance'
    return normalized

def _normalize_word_count(value):
    """Coerce word count into an integer. If a range is given, pick the max."""
    if value is None:
        return None
    # Already numeric
    if isinstance(value, (int, float)):
        return int(value)
    # Extract digits from strings like "3500-4000" and use the maximum
    import re
    numbers = re.findall(r'\d+', str(value))
    if not numbers:
        return None
    return max(int(n) for n in numbers)

def _normalize_level(value):
    """Normalize level to allowed choices."""
    if not value:
        return None
    normalized = str(value).strip().lower()
    mapping = {
        'basic': 'basic',
        'beginner': 'basic',
        'intermediate': 'intermediate',
        'mid': 'intermediate',
        'middle': 'intermediate',
        'advanced': 'advanced',
        'advance': 'advanced',
    }
    return mapping.get(normalized, None)

def _infer_level(word_count=None, instruction_text=None, category=None):
    """Heuristic fallback to guess level when model does not specify."""
    wc = _normalize_word_count(word_count)
    text = (instruction_text or "").lower()
    # Keyword cues
    advanced_keywords = ['phd', 'doctoral', 'masters', 'master', 'dissertation', 'thesis', 'capstone']
    intermediate_keywords = ['case study', 'critical analysis', 'research', 'evaluation']
    if any(k in text for k in advanced_keywords):
        return 'advanced'
    if any(k in text for k in intermediate_keywords):
        return 'intermediate'
    # Word-count heuristic
    if wc is not None:
        if wc >= 4000:
            return 'advanced'
        if wc >= 2000:
            return 'intermediate'
    # Category heuristic (finance/it often mid+)
    if category and str(category).upper() in ('FINANCE', 'IT'):
        return 'intermediate'
    return 'basic'
def role_required(allowed_roles):
    """Decorator to restrict access based on user role"""
    def decorator(view_func):
        def wrapper(request, *args, **kwargs):
            if request.user.role not in allowed_roles:
                messages.error(request, 'You do not have permission to access this page.')
                return redirect('home_dashboard')
            return view_func(request, *args, **kwargs)
        return wrapper
    return decorator


@login_required
@role_required(['marketing'])
def marketing_dashboard(request):
    """Marketing Dashboard - Overview of all marketing activities"""
    
    user = request.user
    
    # Get all jobs created by this marketing user
    all_jobs = Job.objects.filter(created_by=user)
    
    # Get statistics
    stats = {
        'total_jobs': all_jobs.count(),
        'pending_jobs': all_jobs.filter(status='pending').count(),
        'allocated_jobs': all_jobs.filter(status='allocated').count(),
        'completed_jobs': all_jobs.filter(status='completed').count(),
        'hold_jobs': all_jobs.filter(status='hold').count(),
        'query_jobs': all_jobs.filter(status='query').count(),
    }
    
    # Get draft jobs (not yet finalized)
    draft_jobs = all_jobs.filter(status='draft').order_by('-created_at')
    
    # Get recent activities (pending and allocated jobs)
    recent_activities = all_jobs.exclude(status='draft').order_by('-created_at')[:10]
    
    context = {
        'user': user,
        'stats': stats,
        'draft_jobs': draft_jobs,
        'recent_activities': recent_activities,
        'today_date': timezone.now(),
        'status_choices': Job.STATUS_CHOICES,
        'category_choices': Job.CATEGORY_CHOICES,
    }
    
    logger.info(f"Marketing dashboard accessed by: {user.email}")
    return render(request, 'marketing/marketing_dashboard.html', context)
# Event keys for job activities
JOB_EVENTS = {
    'created': 'job.created',
    'initial_saved': 'job.initial_form.saved',
    'initial_submitted': 'job.initial_form.submitted',
    'id_validated': 'job.job_id.validated',
    'summary_requested': 'job.ai_summary.requested',
    'summary_generated': 'job.ai_summary.generated',
    'summary_accepted': 'job.ai_summary.accepted',
    'status_changed': 'job.status.changed',
}


PROXY_ENV_VARS = [
    'OPENAI_PROXY', 'HTTPS_PROXY', 'https_proxy',
    'HTTP_PROXY', 'http_proxy', 'ALL_PROXY', 'all_proxy'
]


# @contextmanager
# def openai_client():
#     """Create OpenAI client after stripping proxy env vars (djongo env sets them globally)."""
#     removed = {}
#     for key in PROXY_ENV_VARS:
#         if key in os.environ:
#             removed[key] = os.environ.pop(key)
    
#     client = None
#     try:
#         client = OpenAI()
#         yield client
#     finally:
#         if client:
#             client.close()
#         for key, value in removed.items():
#             os.environ[key] = value
@contextmanager
def openai_client():
    """OpenAI client using API key only — NO PROXY."""
    client = OpenAI(api_key=settings.OPENAI_API_KEY)
    try:
        yield client
    finally:
        client.close()

def _render_job_list(request, queryset, page_title, filter_description=None,
                     empty_title=None, empty_description=None, template_name='marketing/job_list.html'):
    """Shared renderer for marketing job list pages"""
    paginator = Paginator(queryset, 25)
    page_number = request.GET.get('page')
    jobs_page = paginator.get_page(page_number)
    
    context = {
        'jobs': jobs_page,
        'page_title': page_title,
        'filter_description': filter_description,
        'total_jobs': queryset.count(),
        'empty_state': {
            'title': empty_title or 'No jobs found',
            'description': empty_description or 'Try adjusting the filters or create a new job to get started.',
            'cta_url': reverse('create_job'),
            'cta_label': 'Create Job'
        },
        'status_choices': Job.STATUS_CHOICES,
        'category_choices': Job.CATEGORY_CHOICES,
    }
    return render(request, template_name, context)
def validate_file(file):
    """Validate uploaded file"""
    ALLOWED_EXTENSIONS = ['pdf', 'docx', 'jpg', 'jpeg', 'png']
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
    
    # Get file extension
    ext = os.path.splitext(file.name)[1].lower().replace('.', '')
    
    if ext not in ALLOWED_EXTENSIONS:
        return False, f"File type '.{ext}' not allowed. Allowed types: {', '.join(ALLOWED_EXTENSIONS)}"
    
    if file.size > MAX_FILE_SIZE:
        return False, f"File size exceeds 10MB. Current size: {file.size / (1024*1024):.2f}MB"
    
    return True, "Valid"


@login_required
@role_required(['marketing'])
def create_job(request):
    """Create a new job with two-step form"""
    
    # Check if editing existing draft
    job_id = request.GET.get('job_id')
    job = None
    
    if job_id:
        job = get_object_or_404(Job, system_id=job_id, created_by=request.user, status='draft')
    
    context = {
        'user': request.user,
        'job': job,
    }
    
    return render(request, 'marketing/create_job.html', context)


@login_required
@role_required(['marketing'])
@require_http_methods(["POST"])
def check_job_id_unique(request):
    """AJAX endpoint to check if job_id is unique"""
    try:
        data = json.loads(request.body)
        job_id = data.get('job_id', '').strip()
        current_system_id = data.get('system_id')
        
        if not job_id:
            return JsonResponse({'unique': False, 'message': 'Job ID is required'})
        
        # Check if job_id exists (excluding current job if editing)
        query = Job.objects.filter(job_id=job_id)
        if current_system_id:
            query = query.exclude(system_id=current_system_id)
        
        exists = query.exists()
        
        if exists:
            return JsonResponse({'unique': False, 'message': 'Job ID already exists'})
        else:
            return JsonResponse({'unique': True, 'message': 'Job ID is available'})
            
    except Exception as e:
        logger.error(f"Error checking job ID uniqueness: {str(e)}")
        return JsonResponse({'unique': False, 'message': 'Error checking uniqueness'})


@login_required
@role_required(['marketing'])
@require_http_methods(["POST"])
def save_initial_form(request):
    """Save or update initial form data"""
    try:
        job_id = request.POST.get('job_id', '').strip()
        instruction = request.POST.get('instruction', '').strip()
        files = request.FILES.getlist('attachments')
        system_id = request.POST.get('system_id')  # If editing existing
        remove_ids = request.POST.getlist('remove_attachments')
        replace_flag = request.POST.get('replace_attachments') == 'true'
        
        # Validation
        if not job_id:
            return JsonResponse({'success': False, 'message': 'Job ID is required'}, status=400)
        
        if len(instruction) < 50:
            return JsonResponse({
                'success': False, 
                'message': f'Instruction must be at least 50 characters. Current: {len(instruction)}'
            }, status=400)
        
        # Validate each new file
        if len(files) > 10:
            return JsonResponse({'success': False, 'message': 'Maximum 10 files allowed'}, status=400)
        for file in files:
            is_valid, msg = validate_file(file)
            if not is_valid:
                return JsonResponse({'success': False, 'message': msg}, status=400)

        # Determine attachment counts (existing + new) based on requested operations
        existing_remaining = 0
        if system_id:
            job = Job.objects.get(system_id=system_id, created_by=request.user)
            if replace_flag:
                # Replace wipes all existing, so only new uploads remain
                existing_remaining = 0
            else:
                # Keep everything except the ones explicitly marked for removal
                existing_remaining = job.attachments.exclude(id__in=remove_ids).count()
        else:
            job = None

        total_after = existing_remaining + len(files)
        if total_after <= 0:
            return JsonResponse({'success': False, 'message': 'At least one attachment is required (keep existing or upload new files).'}, status=400)
        if total_after > 10:
            return JsonResponse({'success': False, 'message': 'Maximum 10 attachments allowed including existing and new.'}, status=400)
        
        with transaction.atomic():
            # Create or update job
            if system_id:
                job = Job.objects.filter(system_id=system_id, created_by=request.user).order_by('-created_at').first()
                if not job:
                    return JsonResponse({'success': False, 'message': 'Job not found'}, status=404)
                job.job_id = job_id
                job.instruction = instruction
                job.initial_form_last_saved_at = timezone.now()
                
                # Delete old attachments if replacing
                if replace_flag:
                    job.attachments.all().delete()
                else:
                    if remove_ids:
                        job.attachments.filter(id__in=remove_ids).delete()                
                log_action = 'initial_form_saved'
                event_key = JOB_EVENTS['initial_saved']
            else:
                # Create new job
                system_id = Job.generate_system_id()
                job = Job.objects.create(
                    system_id=system_id,
                    job_id=job_id,
                    instruction=instruction,
                    created_by=request.user,
                    status='draft',
                    job_name_validated_at=timezone.now()
                )
                log_action = 'created'
                event_key = JOB_EVENTS['created']
            
            # Ensure the job is saved before related objects
            job.save()
            job.refresh_from_db()
            
            # Save attachments
            for file in files:
                JobAttachment.objects.create(
                    job=job,
                    file=file,
                    original_filename=file.name,
                    file_size=file.size,
                    uploaded_by=request.user
                )
            
            # Log to JobActionLog
            JobActionLog.objects.create(
                job=job,
                action=log_action,
                performed_by=request.user,
                performed_by_type='user',
                details={
                    'job_id': job_id,
                    'instruction_length': len(instruction),
                    'attachments_count': total_after
                }
            )
            
            # Log to ActivityLog (your system-wide log)
            ActivityLog.objects.create(
                event_key=event_key,
                category='job_management',
                subject_user=request.user,
                performed_by=request.user,
                metadata={
                    'job_system_id': system_id,
                    'job_id': job_id,
                    'instruction_length': len(instruction),
                    'attachments_count': total_after,
                    'status': 'draft'
                }
            )
            
            job.save()
            
            return JsonResponse({
                'success': True,
                'message': 'Initial form saved successfully',
                'system_id': system_id,
                'job_id': job_id
            })
            
    except Job.DoesNotExist:
        return JsonResponse({'success': False, 'message': 'Job not found'}, status=404)
    except Exception as e:
        logger.error(f"Error saving initial form: {str(e)}")
        return JsonResponse({'success': False, 'message': f'Error: {str(e)}'}, status=500)


@login_required
@role_required(['marketing'])
@require_http_methods(["POST"])
def generate_ai_summary(request):
    """Generate AI summary using OpenAI"""
    try:
        data = json.loads(request.body)
        system_id = data.get('system_id')
        
        if not system_id:
            return JsonResponse({'success': False, 'message': 'System ID is required'}, status=400)
        
        job = get_object_or_404(Job, system_id=system_id, created_by=request.user)
        
        # Check if can regenerate
        if not job.can_regenerate_summary():
            return JsonResponse({
                'success': False,
                'message': 'Maximum 3 summary generations reached'
            }, status=400)
        
        # Update timestamps
        job.ai_summary_requested_at = timezone.now()

        # -------------------------------------------------------------------
        # UPDATED: Extract REAL TEXT from PDF, DOCX and Image OCR
        # -------------------------------------------------------------------
        attachments_text = []
        attachments_summary = []
        for attachment in job.attachments.all():
            try:
                ext = attachment.get_file_extension().lower()
                file_path = attachment.file.path
                file_name = attachment.original_filename
			 # -------- PDF Extraction --------
                if ext == ".pdf":
                    try:
                        pdf_text = ""
                        try:
                            from pdfminer.high_level import extract_text
                            pdf_text = extract_text(file_path) or ""
                        except ImportError:
                            # Fallback to PyPDF2 if pdfminer not installed
                            try:
                                import PyPDF2
                                with open(file_path, "rb") as fh:
                                    reader = PyPDF2.PdfReader(fh)
                                    pdf_text = "\n".join(page.extract_text() or "" for page in reader.pages)
                            except Exception as e_fallback:
                                logger.warning(f"PDF fallback extraction error for {file_name}: {e_fallback}")
                                pdf_text = ""
                        if pdf_text.strip():
                            attachments_text.append(
                                f"=== FILE: {file_name} (PDF) ===\n{pdf_text}\n=== END OF {file_name} ===\n"
                            )
                            attachments_summary.append(f"✓ {file_name} (PDF - {len(pdf_text)} chars)")
                        else:
                            attachments_text.append(f"[PDF - {file_name}] (No extractable text)")
                            attachments_summary.append(f"⚠ {file_name} (PDF - No text)")
                    except Exception as e:
                        logger.warning(f"PDF extraction error for {file_name}: {str(e)}")
                        attachments_text.append(f"[PDF - {file_name}] (Extraction failed: {str(e)})")
                        attachments_summary.append(f"✗ {file_name} (PDF - Failed)")
                

				# -------- DOCX Extraction --------
                elif ext == ".docx":
                    try:
                        import docx
                        doc = docx.Document(file_path)
                        
                        # Extract paragraphs
                        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                        
                        # Extract tables
                        tables_text = []
                        for table in doc.tables:
                            for row in table.rows:
                                row_text = ' | '.join([cell.text.strip() for cell in row.cells])
                                if row_text.strip():
                                    tables_text.append(row_text)
                        
                        docx_text = "\n".join(paragraphs)
                        if tables_text:
                            docx_text += "\n\nTABLES:\n" + "\n".join(tables_text)
                        
                        if docx_text.strip():
                            attachments_text.append(
                                f"=== FILE: {file_name} (DOCX) ===\n{docx_text}\n=== END OF {file_name} ===\n"
                            )
                            attachments_summary.append(f"✓ {file_name} (DOCX - {len(docx_text)} chars)")
                        else:
                            attachments_text.append(f"[DOCX - {file_name}] (Empty or unreadable)")
                            attachments_summary.append(f"⚠ {file_name} (DOCX - Empty)")
                    except Exception as e:
                        logger.error(f"DOCX extraction error for {file_name}: {str(e)}")
                        attachments_text.append(f"[DOCX - {file_name}] (Extraction failed: {str(e)})")
                        attachments_summary.append(f"✗ {file_name} (DOCX - Failed)")
                
                # -------- IMAGE OCR --------
                elif ext in ['.png', '.jpg', '.jpeg']:
                    try:
                        from PIL import Image
                        import pytesseract
                        
                        img = Image.open(file_path)
                        # Try OCR with better configuration
                        img_text = pytesseract.image_to_string(
                            img,
                            config='--psm 6'  # Assume uniform block of text
                        )
                        
                        if img_text.strip():
                            attachments_text.append(
                                f"=== FILE: {file_name} (IMAGE OCR) ===\n{img_text}\n=== END OF {file_name} ===\n"
                            )
                            attachments_summary.append(f"✓ {file_name} (Image - {len(img_text)} chars OCR)")
                        else:
                            attachments_text.append(f"[IMAGE - {file_name}] (No readable text)")
                            attachments_summary.append(f"⚠ {file_name} (Image - No text)")
                    except Exception as e:
                        logger.error(f"OCR error for {file_name}: {str(e)}")
                        attachments_text.append(f"[IMAGE - {file_name}] (OCR failed: {str(e)})")
                        attachments_summary.append(f"✗ {file_name} (Image - OCR failed)")
                else:
                    attachments_text.append(f"[Unsupported File Type - {file_name}]")
                    attachments_summary.append(f"⚠ {file_name} (Unsupported type)")
                    
            except Exception as e:
                logger.error(f"Error processing attachment {attachment.original_filename}: {str(e)}")
                attachments_text.append(f"[Error reading {attachment.original_filename}: {str(e)}]")
                attachments_summary.append(f"✗ {attachment.original_filename} (Error)")
        
        # Create attachment summary for logging
        attachment_info = "\n".join(attachments_summary) if attachments_summary else "No attachments processed"
        logger.info(f"Attachment processing summary for {system_id}:\n{attachment_info}")

        # -------------------------------------------------------------------
        # OpenAI Prompt (UNCHANGED BY REQUEST)
        # -------------------------------------------------------------------
        attachments_content = "\n\n".join(attachments_text) if attachments_text else "No attachments or no readable content found in attachments."
        
        prompt = f"""You are an Assignment Analysis Agent for a technical academic writing and student assignment support company. Your role is to carefully analyze the job instruction and any attachments provided.
			CRITICAL INSTRUCTIONS:
			1. Read and analyze ALL provided content thoroughly - both instruction text AND all attachment contents
			2. Extract information from EVERY file provided in the attachments section
			3. If multiple files are provided, synthesize information from ALL of them
			4. Pay special attention to requirements, specifications, learning outcomes, rubrics, and technical details in attachments
			5. If attachments contain assignment briefs, rubrics, or requirements - USE THEM as primary source
            RULE FOR ANALYSIS:
            - If attachments are available, analyze BOTH the instruction and the attachment content.
            - If attachments are NOT available or contain no readable text, analyze ONLY the instruction.
            - Do NOT generate any default summary; base all analysis strictly on the given content.
			TASK: Analyze the job instruction and ALL attachments to extract the following information:
            Your tasks include:
            - Identifying whether the assignment requires the use of any specific software; if yes, specify the exact software name and version (if mentioned or typically required).
            - Providing a detailed task breakdown for any software-related work.
            - Detecting if a PowerPoint presentation is required (e.g., “10-minute presentation”). If yes:
            - Approximate number of slides (default: 1 slide per minute)
            - Estimated words per slide (default: 100 words per slide)
            - Detecting if a LaTeX file is required.
            - Detecting if a poster is required.
            - Estimating the word count if not explicitly mentioned (based on academic standards).
            - Providing a clear structured breakdown of what needs to be written or implemented—without giving the solution itself.

            Use the details below to generate the output.

            INSTRUCTION:
            {job.instruction}

            ATTACHMENTS:
            {attachments_content}

            Generate a detailed JSON response with the following fields:


            1. **Topic**: Extract the exact topic/title from instruction or attachments. If not explicitly stated, create a clear, specific topic based on the content.

			2. **Word Count**: Extract the EXACT word count mentioned in instruction or attachments. If not specified, DO NOT assume - leave as null.And check if there is a limit or range (e.g., "3000-3500 words") and extract the maximum number in that case. Check all attachments for this info. and check all words related to word count like "length", "words", "pages", etc.

			3. **Referencing Style**: Extract the EXACT referencing style mentioned (Harvard, APA, MLA, IEEE, Vancouver, Chicago). If not specified, leave as null. DO NOT assume.

			4. **Writing Style**: Identify the writing style from: proposal, report, essay, dissertation, business_report, personal_development, reflection_writing, case_study. If not specified, leave as null.

			5. **Category**: Determine if this is IT, NON-IT, or FINANCE based on:
			   - IT: Programming, software development, databases, networks, web development, apps, algorithms, data structures, cybersecurity
			   - FINANCE: Accounting, financial analysis, investment, banking, economics, financial management, budgeting
			   - NON-IT: All other subjects (business, marketing, HR, psychology, nursing, etc.)

			6. **Level**: Determine the academic/complexity level using cues from learning outcomes, grading rubrics, and task complexity:
			   - Basic: Undergraduate level 1-2, simple concepts, basic analysis
			   - Intermediate: Undergraduate level 3-4, moderate complexity, good analysis required
			   - Advanced: Masters/PhD level, complex analysis, research-oriented, critical evaluation

			7. **Software**: If ANY specific software, tools, or programming languages are mentioned or required, list them. Examples:
			   - Programming: Python, Java, C++, JavaScript, R, MATLAB etc.
			   - Data Analysis: SPSS, Excel, Tableau, Power BI
			   - Design: AutoCAD, SolidWorks, Adobe Suite
			   - Other: Any specific tools mentioned
			   If NO software is mentioned or required, leave as null.

			8. **Job Summary**: Write a COMPREHENSIVE summary (minimum 250 words) that includes:
			   - Complete overview of what needs to be done
			   - Key requirements from ALL attachments
			   - Structure and format expectations
			   - Software-related tasks (if applicable) - describe what needs to be built/analyzed
			   - Deliverables expected
			   - Any specific guidelines, rubrics, or marking criteria mentioned
			   - Special requirements or constraints
			   - Academic expectations and standards required
   
			   IMPORTANT: Your summary should prove you read ALL attachments by referencing specific details from them.

			CRITICAL RULES:
			- ONLY extract information that is EXPLICITLY stated or clearly evident
			- If something is not mentioned, use null (not a guess)
			- For word count: ONLY use exact numbers stated, never estimate
			- For referencing/writing style: ONLY use exact matches from the allowed options
			- For software: ONLY list if explicitly mentioned or clearly required for the task
			- Be thorough in the summary - this is your chance to show you read everything

			Return ONLY valid JSON in this exact format:
			{{
			    "topic": "extracted or created topic",
			    "word_count": 2000 or null,
			    "referencing_style": "harvard" or null,
			    "writing_style": "report" or null,
			    "category": "IT" or "NON-IT" or "FINANCE",
			    "level": "basic" or "intermediate" or "advanced",
			    "software": ["Python", "Excel"] or null,
			    "job_summary": "comprehensive summary here..."
			}}"""

        # -------------------------------------------------------------------
        # OpenAI CLIENT (UNCHANGED)
        # -------------------------------------------------------------------
        import os
        api_key = os.getenv('OPENAI_API_KEY')
        
        if not api_key:
            return JsonResponse({
                'success': False,
                'message': 'OpenAI API key not configured'
            }, status=500)
        
        from openai import OpenAI
        client = OpenAI(
            api_key=api_key,
            timeout=90.0,
            max_retries=2
        )
        
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {
                    "role": "system", 
                    "content": "You are an expert academic assignment analyzer. You MUST read and analyze ALL provided content including every attachment. Always return valid JSON. Be thorough and accurate."
                },
                {
                    "role": "user", 
                    "content": prompt
                }
            ],
            temperature=0.3,
            max_tokens=3000
        )
        
        ai_response = response.choices[0].message.content.strip()
        
        # Clean JSON
        if ai_response.startswith('```'):
            ai_response = ai_response.split('```')[1]
            if ai_response.startswith('json'):
                ai_response = ai_response[4:]
            ai_response = ai_response.strip()
        
        summary_data = json.loads(ai_response)

        # -------------------------------------------------------------------
        # SAVE SUMMARY (UNCHANGED)
        # -------------------------------------------------------------------
        with transaction.atomic():
            job.topic = summary_data.get('topic')
            job.word_count = _normalize_word_count(summary_data.get('word_count'))
            job.referencing_style = summary_data.get('referencing_style')
            job.writing_style = summary_data.get('writing_style')
            job.job_summary = summary_data.get('job_summary')
            job.category = summary_data.get('category')
            
            # Store level and software in job_summary metadata or as JSON field
            level = _normalize_level(summary_data.get('level'))
            software_list = summary_data.get('software')
            
            # Format software as string
            if software_list and isinstance(software_list, list):
                job.software = ", ".join(software_list)
            else:
                job.software = None
            job.level = level or _infer_level(job.word_count, job.instruction, job.category)
            
            # Increment version
            job.ai_summary_version += 1
            
            generation_timestamps = job.ai_summary_generated_at or []
            generation_timestamps.append(timezone.now().isoformat())
            job.ai_summary_generated_at = generation_timestamps
            
            degree = job.calculate_degree()
            
            JobSummaryVersion.objects.create(
                job=job,
                version_number=job.ai_summary_version,
                topic=job.topic,
                word_count=job.word_count,
                referencing_style=job.referencing_style,
                writing_style=job.writing_style,
                job_summary=job.job_summary,
                degree=degree,
                performed_by='system',
                ai_model_used='gpt-4o'
            )
            
            JobActionLog.objects.create(
                job=job,
                action='ai_summary_generated',
                performed_by=request.user,
                performed_by_type='system',
                details={
                    'version': job.ai_summary_version,
                    'degree': degree,
                    'model': 'gpt-4o',
                    'category': job.category,
                    'level': level,
                    'software': job.software,
                    'attachments_processed': len(attachments_text)
                }
            )

            # -----------------------------
            # AUTO ACCEPT & REDIRECT LOGIC
            # -----------------------------
            auto_accept = False
            auto_redirect = False

            # Rule 1: Perfect summary (degree 0) → Auto accept AND redirect
            if degree == 0:
                auto_accept = True
                auto_redirect = True

            # Rule 2: Version 3 reached → Auto accept but NO redirect
            elif job.ai_summary_version >= 3:
                auto_accept = True
                auto_redirect = False

            # Apply auto-accept
            if auto_accept:
                job.ai_summary_accepted_at = timezone.now()
                job.status = "pending"

                JobActionLog.objects.create(
                    job=job,
                    action="ai_summary_accepted",
                    performed_by=request.user,
                    performed_by_type="system",
                    details={
                        "version": job.ai_summary_version,
                        "degree": degree,
                        "auto_accepted": True,
                        "redirect": auto_redirect
                    }
                )

            
            # auto_accept = job.should_auto_accept()
            # if auto_accept:
            #     job.ai_summary_accepted_at = timezone.now()
            #     job.status = 'pending'
                
            #     JobActionLog.objects.create(
            #         job=job,
            #         action='ai_summary_accepted',
            #         performed_by=request.user,
            #         performed_by_type='system',
            #         details={
            #             'auto_accepted': True,
            #             'reason': 'degree_0' if degree == 0 else 'version_3'
            #         }
            #     )
                
                ActivityLog.objects.create(
                    event_key=JOB_EVENTS['summary_generated'],
                    category='job_management',
                    subject_user=request.user,
                    performed_by=request.user,
                    metadata={
                        'job_system_id': job.system_id,
                        'job_id': job.job_id,
                        'version': job.ai_summary_version,
                        'degree': degree,
                        'model': 'gpt-4o',
                        'attachments_processed': len(attachments_text)
                    }
                )
            
            job.save()
            
            return JsonResponse({
                'success': True,
                'message': 'AI summary generated successfully',
                'data': {
                    'topic': job.topic,
                    'word_count': job.word_count,
                    'referencing_style': job.referencing_style,
                    'writing_style': job.writing_style,
                    'job_summary': job.job_summary,
					'category': job.category,
                    'level': level,
                    'software': job.software,
                    'version': job.ai_summary_version,
                    'degree': degree,
                    'auto_accepted': auto_accept,
                    'auto_redirect': auto_redirect,
                    'can_regenerate': job.can_regenerate_summary(),
					'attachments_processed': len(attachments_text)
                }
            })

    except Job.DoesNotExist:
        return JsonResponse({'success': False, 'message': 'Job not found'}, status=404)
    except json.JSONDecodeError as e:
        logger.error(f"JSON decode error: {str(e)}\nAI Response: {ai_response}")
        return JsonResponse({'success': False, 'message': 'Error parsing AI response'}, status=500)
    except Exception as e:
        logger.error(f"Error generating AI summary: {str(e)}")
        return JsonResponse({'success': False, 'message': f'Error: {str(e)}'}, status=500)




@login_required
@role_required(['marketing'])
@require_http_methods(["POST"])
def accept_summary(request):
    """Accept AI summary and finalize job"""
    try:
        data = json.loads(request.body)
        system_id = data.get('system_id')
        
        job = get_object_or_404(Job, system_id=system_id, created_by=request.user)
        
        with transaction.atomic():
            job.ai_summary_accepted_at = timezone.now()
            job.status = 'pending'
            job.initial_form_submitted_at = timezone.now()
            
            # Log to JobActionLog
            JobActionLog.objects.create(
                job=job,
                action='ai_summary_accepted',
                performed_by=request.user,
                performed_by_type='user',
                details={
                    'version': job.ai_summary_version,
                    'degree': job.job_card_degree
                }
            )
            
            JobActionLog.objects.create(
                job=job,
                action='initial_form_submitted',
                performed_by=request.user,
                performed_by_type='user'
            )
            
            # Log to ActivityLog
            ActivityLog.objects.create(
                event_key=JOB_EVENTS['summary_accepted'],
                category='job_management',
                subject_user=request.user,
                performed_by=request.user,
                metadata={
                    'job_system_id': job.system_id,
                    'job_id': job.job_id,
                    'version': job.ai_summary_version,
                    'degree': job.job_card_degree,
                    'manual_acceptance': True
                }
            )
            
            ActivityLog.objects.create(
                event_key=JOB_EVENTS['initial_submitted'],
                category='job_management',
                subject_user=request.user,
                performed_by=request.user,
                metadata={
                    'job_system_id': job.system_id,
                    'job_id': job.job_id,
                    'status': 'pending'
                }
            )
            
            job.save()
            
            return JsonResponse({
                'success': True,
                'message': 'Job created successfully',
                'redirect': '/marketing/dashboard/'
            })
            
    except Job.DoesNotExist:
        return JsonResponse({'success': False, 'message': 'Job not found'}, status=404)
    except Exception as e:
        logger.error(f"Error accepting summary: {str(e)}")
        return JsonResponse({'success': False, 'message': f'Error: {str(e)}'}, status=500)


@login_required
@role_required(['marketing'])
def get_summary_versions(request, system_id):
    """Get all summary versions for a job"""
    try:
        job = get_object_or_404(Job, system_id=system_id, created_by=request.user)
        versions = job.summary_versions.all()
        
        versions_data = [{
            'version': v.version_number,
            'topic': v.topic,
            'word_count': v.word_count,
            'referencing_style': v.referencing_style,
            'writing_style': v.writing_style,
            'job_summary': v.job_summary,
            'degree': v.degree,
            'generated_at': v.generated_at.isoformat()
        } for v in versions]
        
        return JsonResponse({
            'success': True,
            'versions': versions_data,
            'current_version': job.ai_summary_version
        })
        
    except Job.DoesNotExist:
        return JsonResponse({'success': False, 'message': 'Job not found'}, status=404)
    except Exception as e:
        return JsonResponse({'success': False, 'message': str(e)}, status=500)


@login_required
@role_required(['marketing'])
def my_jobs(request):
    """List all jobs created by the current marketing user"""
    queryset = Job.objects.filter(
        created_by=request.user
    ).select_related('allocated_to').order_by('-created_at')
    
    return _render_job_list(
        request,
        queryset,
        page_title='My Jobs',
        filter_description='Full history of every job you have created.',
        empty_title='No jobs yet',
        empty_description='Create a job to see it listed here.'
    )


@login_required
@role_required(['marketing'])
def hold_jobs(request):
    """Jobs currently on hold"""
    queryset = Job.objects.filter(
        created_by=request.user,
        status='hold'
    ).select_related('allocated_to').order_by('-updated_at')
    
    return _render_job_list(
        request,
        queryset,
        page_title='Hold Jobs',
        filter_description='Jobs paused for clarification or awaiting client confirmation.',
        empty_title='No jobs on hold',
        empty_description='When you pause a job it will surface here for quick follow-up.'
    )


@login_required
@role_required(['marketing'])
def query_jobs(request):
    """Jobs flagged with queries"""
    queryset = Job.objects.filter(
        created_by=request.user,
        status='query'
    ).select_related('allocated_to').order_by('-updated_at')
    
    return _render_job_list(
        request,
        queryset,
        page_title='Query Jobs',
        filter_description='Jobs that need action because the allocator or writer raised queries.',
        empty_title='No query jobs',
        empty_description='Great news—no active queries right now.'
    )


@login_required
@role_required(['marketing'])
def unallocated_jobs(request):
    """Jobs that have not been allocated yet"""
    queryset = Job.objects.filter(
        created_by=request.user,
        allocated_to__isnull=True
    ).exclude(
        status__in=['draft', 'completed', 'cancelled']
    ).order_by('-created_at')
    
    return _render_job_list(
        request,
        queryset,
        page_title='Unallocated Jobs',
        filter_description='Submitted jobs still waiting for allocator assignment.',
        empty_title='No unallocated jobs',
        empty_description='All submitted jobs have already been assigned.'
    )


@login_required
@role_required(['marketing'])
def completed_jobs(request):
    """Jobs completed by delivery teams"""
    queryset = Job.objects.filter(
        created_by=request.user,
        status='completed'
    ).select_related('allocated_to').order_by('-updated_at')
    
    return _render_job_list(
        request,
        queryset,
        page_title='Completed Jobs',
        filter_description='Finished jobs delivered back by the production teams.',
        empty_title='No completed jobs yet',
        empty_description='Once a job is delivered successfully, it will be archived here.'
    )


@login_required
@role_required(['marketing'])
def allocated_jobs(request):
    """Jobs that are currently allocated to a downstream team"""
    queryset = Job.objects.filter(
        created_by=request.user,
        status__in=['allocated', 'in_progress']
    ).select_related('allocated_to').order_by('-updated_at')
    
    return _render_job_list(
        request,
        queryset,
        page_title='Allocated Jobs',
        filter_description='Live jobs currently being worked on by writers or process teams.',
        empty_title='No allocated jobs',
        empty_description='Jobs will appear here as soon as the allocator assigns them.'
    )





@login_required
@role_required(['marketing'])
def final_job_form(request, system_id):
    """Final Job Drop Form - Two column layout"""
    
    job = get_object_or_404(Job, system_id=system_id, created_by=request.user)
    
    # Check if job status allows final form
    if job.status not in ['draft', 'pending']:
        messages.warning(request, 'This job has already been finalized.')
        return redirect('marketing_dashboard')
    
    # Log that final form was opened
    if not job.final_form_opened_at:
        job.final_form_opened_at = timezone.now()
        job.save(update_fields=['final_form_opened_at'])
        
        log_activity_event(
            'job.final_form_opened_at',
            subject_user=request.user,
            performed_by=request.user,
            metadata={
                'job_system_id': job.system_id,
                'job_id': job.job_id,
            },
        )
    
    if request.method == 'POST':
        return _process_final_form_submission(request, job)
    
    # Get form options
    templates = [
        template for template in TemplateMaster.objects.filter(status='active')
        if not getattr(template, 'is_deleted', False)
    ]
    project_groups = [
        group for group in ProjectGroupMaster.objects.all()
        if not getattr(group, 'is_deleted', False)
    ]
    referencing_styles = [
        ref for ref in ReferencingMaster.objects.all()
        if not getattr(ref, 'is_deleted', False)
    ]
    writing_styles = [
        writing for writing in AcademicWritingMaster.objects.all()
        if not getattr(writing, 'is_deleted', False)
    ]
    price_entries = [
        entry for entry in PriceMaster.objects.all()
        if not getattr(entry, 'is_deleted', False)
    ]
    price_map = {}
    for entry in price_entries:
        category_key = str(entry.category).upper()
        level_key = str(entry.level).lower()
        level_label = entry.get_level_display() if hasattr(entry, 'get_level_display') else level_key.title()
        price_map.setdefault(category_key, {})[level_key] = {
            'price': _to_float(entry.price_per_word),
            'label': level_label,
        }
    selected_price_level = _to_price_master_level(job.level)
        
    # Get existing attachments
    existing_attachments = job.attachments.all()
    customers_qs = Customer.objects.filter(created_by=request.user).order_by('customer_name')
    customers = [c for c in customers_qs if getattr(c, 'is_active', False)]
    
    context = {
        'job': job,
        'templates': templates,
        'project_groups': project_groups,
        'referencing_styles': referencing_styles,
        'writing_styles': writing_styles,
        'existing_attachments': existing_attachments,
        'category_choices': Job.CATEGORY_CHOICES,
        'referencing_choices': Job.REFERENCING_STYLE_CHOICES,
        'writing_choices': Job.WRITING_STYLE_CHOICES,
        'price_map': json.dumps(price_map),
        'selected_price_level': selected_price_level or '',
        'customers': customers,
    }
    
    return render(request, 'marketing/final_job_form.html', context)


def _process_final_form_submission(request, job):
    """Process final form submission"""
    try:
        # Extract form data
        topic = request.POST.get('topic', '').strip() or job.topic
        category = request.POST.get('category', '').strip()
        level = request.POST.get('level', '').strip()
        word_count = request.POST.get('word_count', '').strip()
        referencing_style = request.POST.get('referencing_style', '').strip()
        writing_style = request.POST.get('writing_style', '').strip()
        instruction = request.POST.get('instruction', '').strip()
        template_id = request.POST.get('template', '').strip()
        project_group_id = request.POST.get('project_group', '').strip()
        expected_deadline = request.POST.get('expected_deadline', '').strip()
        strict_deadline = request.POST.get('strict_deadline', '').strip()
        software = request.POST.get('software', '').strip()
        amount = request.POST.get('amount', '').strip()
        customer_id = request.POST.get('customer_id', '').strip()
        
        # Validation
        errors = []
        
        if not category:
            errors.append('Category is required.')
        
        if not level:
            errors.append('Level is required.')
        
        if not word_count or not word_count.isdigit():
            errors.append('Valid word count is required.')
        else:
            word_count = int(word_count)
            if word_count <= 0:
                errors.append('Word count must be greater than 0.')
        
        if not instruction or len(instruction) < 50:
            errors.append('Instruction must be at least 50 characters.')
        
        if not template_id:
            errors.append('Template selection is required.')
        
        if not project_group_id:
            errors.append('Project group selection is required.')
        
        if not expected_deadline:
            errors.append('Expected deadline is required.')
        
        if not strict_deadline:
            errors.append('Strict deadline is required.')
        
        if not customer_id:
            errors.append('Customer selection is required.')
        
        if not amount:
            errors.append('Amount is required.')
        else:
            try:
                amount = float(amount)
                if amount <= 0:
                    errors.append('Amount must be greater than 0.')
            except ValueError:
                errors.append('Invalid amount format.')
        
        # Validate deadlines
        if expected_deadline and strict_deadline:
            try:
                expected_dt = timezone.datetime.fromisoformat(expected_deadline)
                strict_dt = timezone.datetime.fromisoformat(strict_deadline)
                
                # Make timezone-aware if needed
                if timezone.is_naive(expected_dt):
                    expected_dt = timezone.make_aware(expected_dt)
                if timezone.is_naive(strict_dt):
                    strict_dt = timezone.make_aware(strict_dt)
                
                # Strict deadline must be at least 24 hours from now
                min_strict = timezone.now() + timedelta(hours=24)
                if strict_dt < min_strict:
                    errors.append('Strict deadline must be at least 24 hours from now.')
                
                # Expected should be before strict
                if expected_dt >= strict_dt:
                    errors.append('Expected deadline should be before strict deadline.')
                
            except ValueError:
                errors.append('Invalid deadline format.')
        
        if errors:
            for error in errors:
                messages.error(request, error)
            return redirect('final_job_form', system_id=job.system_id)
        
        # Resolve customer
        customer_obj = get_object_or_404(
            Customer,
            customer_id=customer_id,
            created_by=request.user
        )
        
        # Get related objects with manual soft-delete checks (djongo friendly)
        template = get_object_or_404(TemplateMaster, id=template_id)
        if getattr(template, 'is_deleted', False):
            messages.error(request, 'Selected template is no longer available.')
            return redirect('final_job_form', system_id=job.system_id)

        project_group = get_object_or_404(ProjectGroupMaster, id=project_group_id)
        if getattr(project_group, 'is_deleted', False):
            messages.error(request, 'Selected project group is no longer available.')
            return redirect('final_job_form', system_id=job.system_id)
        
        # Calculate system expected amount
        system_expected = None
        normalized_category = category.upper()
        price_level = _to_price_master_level(level)
        price_entry = next(
            (
                item for item in PriceMaster.objects.filter(
                    category=normalized_category,
                    level=price_level
                )
                if not getattr(item, 'is_deleted', False)
            ),
            None
        )
        if not price_entry:
            messages.error(request, 'Pricing not configured for the selected category and level.')
            return redirect('final_job_form', system_id=job.system_id)
        
        price_per_word = _to_float(price_entry.price_per_word)
        system_expected = price_per_word * float(word_count)
        normalized_level = _normalize_level(level)
        
        with transaction.atomic():
            # Update job with final form data
            job.topic = topic
            job.category = category
            job.level = normalized_level
            job.word_count = word_count
            job.referencing_style = referencing_style or None
            job.writing_style = writing_style or None
            job.instruction = instruction
            job.template = template
            job.project_group = project_group
            job.expected_deadline = expected_dt
            job.strict_deadline = strict_dt
            job.software = software or None
            job.amount = amount
            job.system_expected_amount = system_expected
            job.customer_name = customer_obj.customer_name
            job.customer_id = customer_obj.customer_id
            
            # Update timestamps
            job.final_form_submitted_at = timezone.now()
            
            # Change status to unallocated
            job.status = 'unallocated'
            
            job.save()
            
            # Handle additional attachments
            new_files = request.FILES.getlist('other_attachments')
            for file in new_files:
                is_valid, msg = validate_file(file)
                if is_valid:
                    JobAttachment.objects.create(
                        job=job,
                        file=file,
                        original_filename=file.name,
                        file_size=file.size,
                        uploaded_by=request.user
                    )
            
            # Generate project prefix
            project_prefix = JobTemplate.generate_project_prefix(project_group)
            
            # Generate masking ID
            masking_id = JobTemplate.generate_masking_id(category, project_prefix)
            masking_generated_at = timezone.now()
            
            # Create JobTemplate instance
            job_template = JobTemplate.objects.create(
                job=job,    
                template=template,
                masking_id=masking_id,
                project_group=project_group,
                project_prefix=project_prefix,
                created_by=request.user,
                masking_id_generated_at=masking_generated_at
            )
            
            # Update job with masking timestamp
            job.masking_id_generated_at = masking_generated_at
            job.save(update_fields=['masking_id_generated_at'])
            
            # Create default tasks
            task_configs = template.default_tasks
            for task_config in task_configs:
                task_id = JobTask.generate_task_id(
                    project_prefix,
                    task_config['task_code']
                )
                
                JobTask.objects.create(
                    job_template=job_template,
                    task_id=task_id,
                    task_number=task_config['task_number'],
                    task_name=task_config['task_name'],
                    task_code=task_config['task_code'],
                    assignable_roles=task_config['assignable_roles'],
                    status='pending',
                    word_count=word_count if task_config['task_number'] == 1 else None
                )
            
            # Log activities
            log_activity_event(
                'job.final_form_submitted_at',
                subject_user=request.user,
                performed_by=request.user,
                metadata={
                    'job_system_id': job.system_id,
                    'job_id': job.job_id,
                    'masking_id': masking_id,
                    'status': 'unallocated',
                },
            )
            
            log_activity_event(
                'job.masking_id_generated_at',
                subject_user=None,
                performed_by=None,
                metadata={
                    'job_system_id': job.system_id,
                    'masking_id': masking_id,
                    'performed_by': 'system',
                },
            )
            
            JobActionLog.objects.create(
                job=job,
                action='final_form_submitted',
                performed_by=request.user,
                performed_by_type='user',
                details={
                    'system_id': job.system_id,
                    'template': template.template_name,
                    'project_group': project_group.project_group_name,
                    'category': category,
                    'level': normalized_level,
                }
            )
        
        messages.success(
            request,
            f'Job "{job.job_id}" has been successfully finalized! '
            f'System ID: {job.system_id}'
        )
        return redirect('marketing_dashboard')
        
    except Exception as e:
        logger.exception(f"Error processing final form: {str(e)}")
        messages.error(request, 'An error occurred while submitting the form.')
        return redirect('final_job_form', system_id=job.system_id)


@login_required
@role_required(['marketing'])
def get_system_expected_amount(request):
    """Return system expected amount based on category and word count"""
    category = request.GET.get('category', '').strip()
    level = request.GET.get('level', '').strip()
    word_count = request.GET.get('word_count', '').strip()

    if not category or not level or not word_count or not word_count.isdigit():
        return JsonResponse({'success': False, 'message': 'Invalid inputs'}, status=400)

    word_count = int(word_count)
    if word_count <= 0:
        return JsonResponse({'success': False, 'message': 'Invalid word count'}, status=400)

    normalized_category = category.upper()
    price_level = _to_price_master_level(level)
    price_entry = next(
        (
            item for item in PriceMaster.objects.filter(
                category=normalized_category,
                level=price_level
            )
            if not getattr(item, 'is_deleted', False)
        ),
        None
    )

    if not price_entry:
        return JsonResponse({'success': False, 'message': 'Pricing not configured'}, status=404)

    # Use _to_float here too
    price_per_word = _to_float(price_entry.price_per_word)
    amount = price_per_word * word_count
    return JsonResponse({
        'success': True,
        'price_per_word': price_per_word,
        'amount': amount,
    })


@login_required
@role_required(['marketing'])
@require_http_methods(["POST"])
def copy_summary_to_final(request):
    """AJAX endpoint to copy summary data to final form"""
    try:
        data = json.loads(request.body)
        system_id = data.get('system_id')
        
        job = get_object_or_404(Job, system_id=system_id, created_by=request.user)
        
        return JsonResponse({
            'success': True,
            'data': {
                'topic': job.topic or '',
                'word_count': job.word_count or '',
                'referencing_style': job.referencing_style or '',
                'writing_style': job.writing_style or '',
                'instruction': job.instruction or '',
            }
        })
        
    except Exception as e:
        logger.error(f"Error copying summary: {str(e)}")
        return JsonResponse({
            'success': False,
            'message': str(e)
        }, status=500)





@login_required
@role_required(['marketing'])
def view_job_details(request, system_id):
    """View detailed job information"""
    
    job = get_object_or_404(Job, system_id=system_id, created_by=request.user)
    
    # Get job template and tasks if exists
    job_template = None
    tasks = []
    if hasattr(job, 'job_template'):
        job_template = job.job_template
        tasks = job_template.tasks.all().order_by('task_number')
    
    # Get all attachments with existence check (avoid broken links)
    attachments = job.attachments.all()
    attachments_display = []
    for att in attachments:
        exists = False
        url = None
        try:
            exists = att.file and att.file.storage.exists(att.file.name)
            if exists:
                url = att.file.url
        except Exception as e:
            logger.warning(f"Attachment missing or unreadable ({att.original_filename}): {e}")
            exists = False
            url = None
        attachments_display.append({
            'obj': att,
            'exists': exists,
            'url': url,
        })
    
    # Get all summary versions
    summary_versions = job.summary_versions.all().order_by('version_number')
    has_ai_summary = (
        summary_versions.exists()
        or bool(job.ai_summary_version)
        or bool(job.job_summary)
        or bool(job.ai_summary_generated_at)
    )
    
    # Get action logs for timeline
    action_logs = job.action_logs.all().order_by('timestamp')
    
    # Build timeline events
    timeline_events = []
    
    # Job Created
    if job.created_at:
        timeline_events.append({
            'timestamp': job.created_at,
            'title': 'Job Created',
            'description': f'System ID: {job.system_id} | Job ID: {job.job_id}',
            'icon': 'plus-circle',
            'color': 'blue'
        })
    
    # Job Name Validated
    if job.job_name_validated_at:
        timeline_events.append({
            'timestamp': job.job_name_validated_at,
            'title': 'Job ID Validated',
            'description': 'Job ID uniqueness confirmed',
            'icon': 'check-circle',
            'color': 'green'
        })
    
    # Initial Form Submitted
    if job.initial_form_submitted_at:
        timeline_events.append({
            'timestamp': job.initial_form_submitted_at,
            'title': 'Initial Form Submitted',
            'description': f'Instruction and {attachments.count()} attachment(s)',
            'icon': 'file-text',
            'color': 'blue'
        })
    
    # AI Summary Requested
    if job.ai_summary_requested_at:
        timeline_events.append({
            'timestamp': job.ai_summary_requested_at,
            'title': 'AI Summary Requested',
            'description': 'Summary generation initiated',
            'icon': 'cpu',
            'color': 'purple'
        })
    
    # AI Summary Generated (all versions)
    if job.ai_summary_generated_at:
        for idx, timestamp_str in enumerate(job.ai_summary_generated_at, 1):
            try:
                timestamp = timezone.datetime.fromisoformat(timestamp_str)
                if timezone.is_naive(timestamp):
                    timestamp = timezone.make_aware(timestamp)
                
                timeline_events.append({
                    'timestamp': timestamp,
                    'title': f'AI Summary Generated (v{idx})',
                    'description': f'Degree: {summary_versions[idx-1].degree if idx <= summary_versions.count() else "N/A"}',
                    'icon': 'zap',
                    'color': 'yellow'
                })
            except Exception as e:
                logger.error(f"Error parsing timestamp: {e}")
    
    # AI Summary Accepted
    if job.ai_summary_accepted_at:
        timeline_events.append({
            'timestamp': job.ai_summary_accepted_at,
            'title': 'AI Summary Accepted',
            'description': f'Version {job.ai_summary_version} accepted',
            'icon': 'check',
            'color': 'green'
        })
    
    # Final Form Opened
    if job.final_form_opened_at:
        timeline_events.append({
            'timestamp': job.final_form_opened_at,
            'title': 'Final Form Opened',
            'description': 'Marketing user accessed final form',
            'icon': 'eye',
            'color': 'gray'
        })
    
    # Final Form Submitted
    if job.final_form_submitted_at:
        timeline_events.append({
            'timestamp': job.final_form_submitted_at,
            'title': 'Final Form Submitted',
             'description': f'Status changed to {job.get_status_display()} | System ID: {job.system_id}',
            'icon': 'send',
            'color': 'blue'
        })
    
    # Masking ID Generated
    if job.masking_id_generated_at and job_template:
        timeline_events.append({
            'timestamp': job.masking_id_generated_at,
            'title': 'Masking ID Generated',
            'description': f'ID: {job_template.masking_id}',
            'icon': 'hash',
            'color': 'purple'
        })
    
    # Sort timeline by timestamp
    timeline_events.sort(key=lambda x: x['timestamp'])

    def _format_currency(value):
        amount = _decimal_to_float(value)
        if amount is None:
            return None
        return f"{amount:,.2f}"

    context = {
        'job': job,
        'job_template': job_template,
        'tasks': tasks,
        'attachments': attachments,
        'attachments_display': attachments_display,
        'summary_versions': summary_versions,
        'timeline_events': timeline_events,
        'job_amount_display': _format_currency(job.amount),
        'job_system_amount_display': _format_currency(job.system_expected_amount),
        'has_ai_summary': has_ai_summary,
    }
    
    return render(request, 'marketing/view_job_details.html', context)

@login_required
@role_required(['marketing'])
def customer_management(request):
    """Customer Management Page - List and Add Customers"""
    
    # Get all customers created by this marketing user
    customers_qs = Customer.objects.filter(
        created_by=request.user
    ).order_by('-created_at')
    # Filter out malformed/partial records so we don't render "None" rows
    customers = [
        customer for customer in customers_qs
        if customer.customer_id and customer.customer_name and customer.customer_email and customer.customer_phone
    ]
    customers_display = []
    for customer in customers:
        customers_display.append({
            'customer_id': customer.customer_id,
            'customer_name': customer.customer_name,
            'customer_email': customer.customer_email,
            'customer_phone': customer.customer_phone,
            'is_active': customer.is_active,
            'targeted_amount': _decimal_to_float(customer.targeted_amount) or 0.0,
            'current_amount': _decimal_to_float(customer.current_amount) or 0.0,
        })
    total_customers = len(customers)
    active_customers = sum(1 for customer in customers if customer.is_active)
    inactive_customers = total_customers - active_customers
    
    context = {
        'customers': customers_display,
        'total_customers': total_customers,
        'active_customers': active_customers,
        'inactive_customers': inactive_customers,
    }
    
    return render(request, 'marketing/customer_management.html', context)


@login_required
@role_required(['marketing'])
@require_http_methods(["POST"])
def add_customer(request):
    """AJAX endpoint to add a new customer"""
    try:
        data = json.loads(request.body)
        
        customer_name = data.get('customer_name', '').strip()
        customer_email = data.get('customer_email', '').strip()
        customer_phone = data.get('customer_phone', '').strip()
        targeted_amount = data.get('targeted_amount', '').strip()
        from decimal import Decimal, InvalidOperation
        
        # Validation
        errors = {}
        
        if not customer_name or len(customer_name) < 3:
            errors['customer_name'] = 'Customer name must be at least 3 characters.'
        
        if not customer_email:
            errors['customer_email'] = 'Email is required.'
        elif Customer.objects.filter(customer_email=customer_email).exists():
            errors['customer_email'] = 'This email is already registered.'
        
        if not customer_phone:
            errors['customer_phone'] = 'Phone number is required.'
        elif not customer_phone.isdigit() or len(customer_phone) != 10:
            errors['customer_phone'] = 'Phone number must be exactly 10 digits.'
        
        if not targeted_amount:
            errors['targeted_amount'] = 'Targeted amount is required.'
        else:
            try:
                targeted_amount = Decimal(str(targeted_amount))
                if targeted_amount < Decimal('1'):
                    errors['targeted_amount'] = 'Targeted amount must be at least 1.'
            except (InvalidOperation, ValueError):
                errors['targeted_amount'] = 'Invalid amount format.'
        
        if errors:
            return JsonResponse({
                'success': False,
                'errors': errors
            }, status=400)
        
        # Create customer
        with transaction.atomic():
            customer = Customer.objects.create(
                customer_name=customer_name,
                customer_email=customer_email,
                customer_phone=customer_phone,
                targeted_amount=targeted_amount,
                created_by=request.user,
                is_active=True
            )
            
            # Log action
            CustomerActionLog.objects.create(
                customer=customer,
                action='created',
                performed_by=request.user,
                details={
                    'customer_name': customer_name,
                    'customer_email': customer_email,
                    'targeted_amount': float(targeted_amount)
                }
            )
            
            # System-wide activity log
            ActivityLog.objects.create(
                event_key='customer.created',
                category='customer_management',
                subject_user=request.user,
                performed_by=request.user,
                metadata={
                    'customer_id': customer.customer_id,
                    'customer_name': customer_name,
                    'customer_email': customer_email,
                }
            )
            
            return JsonResponse({
                'success': True,
                'message': 'Customer added successfully!',
                'customer': {
                    'customer_id': customer.customer_id,
                    'customer_name': customer.customer_name,
                    'customer_email': customer.customer_email,
                    'customer_phone': customer.customer_phone,
                    'targeted_amount': float(customer.targeted_amount or 0),
                    'current_amount': float(customer.current_amount or 0),
                    'is_active': customer.is_active,
                    'created_at': customer.created_at.isoformat(),
                }
            })
            
    except Exception as e:
        logger.error(f"Error adding customer: {str(e)}")
        return JsonResponse({
            'success': False,
            'message': f'Error: {str(e)}'
        }, status=500)


@login_required
@role_required(['marketing'])
@require_http_methods(["POST"])
def toggle_customer_status(request):
    """AJAX endpoint to toggle customer active status"""
    try:
        data = json.loads(request.body)
        customer_id = data.get('customer_id')
        
        customer = get_object_or_404(
            Customer,
            customer_id=customer_id,
            created_by=request.user
        )
        
        # Toggle status without invoking model full_clean (avoids decimal validation noise)
        new_status = not customer.is_active
        Customer.objects.filter(pk=customer.pk).update(
            is_active=new_status,
            updated_at=timezone.now()
        )
        customer.is_active = new_status
        
        # Log action
        action = 'activated' if customer.is_active else 'deactivated'
        CustomerActionLog.objects.create(
            customer=customer,
            action=action,
            performed_by=request.user,
            details={'is_active': customer.is_active}
        )
        
        ActivityLog.objects.create(
            event_key=f'customer.{action}',
            category='customer_management',
            subject_user=request.user,
            performed_by=request.user,
            metadata={
                'customer_id': customer.customer_id,
                'customer_name': customer.customer_name,
                'is_active': customer.is_active,
            }
        )
        
        return JsonResponse({
            'success': True,
            'message': f'Customer {action} successfully!',
            'is_active': customer.is_active
        })
        
    except Exception as e:
        logger.error(f"Error toggling customer status: {str(e)}")
        return JsonResponse({
            'success': False,
            'message': str(e)
        }, status=500)


@login_required
@role_required(['marketing'])
def get_customer_kpis(request, customer_id):
    """AJAX endpoint to get customer KPIs"""
    try:
        customer = get_object_or_404(
            Customer,
            customer_id=customer_id,
            created_by=request.user
        )
        
        # Update KPIs before returning; tolerate djongo query quirks
        try:
            customer.update_kpis()
            customer.refresh_from_db()
        except Exception as kpi_err:
            logger.warning(f"Could not recompute KPIs for customer {customer.customer_id}: {kpi_err}")
        
        return JsonResponse({
            'success': True,
            'kpis': {
                'project_kpis': {
                    'total_projects': customer.total_projects,
                    'completed_projects': customer.completed_projects,
                    'cancelled_projects': customer.cancelled_projects,
                    'projects_with_issues': customer.projects_with_issues,
                },
                'financial_kpis': {
                    'total_order_amount': _decimal_to_float(customer.total_order_amount) or 0.0,
                    'total_paid_amount': _decimal_to_float(customer.total_paid_amount) or 0.0,
                    'remaining_amount': _decimal_to_float(customer.remaining_amount) or 0.0,
                },
                'category_breakdown': {
                    'it_projects': customer.it_projects,
                    'non_it_projects': customer.non_it_projects,
                    'finance_projects': customer.finance_projects,
                }
            }
        })
        
    except Exception as e:
        logger.error(f"Error fetching customer KPIs: {str(e)}")
        return JsonResponse({
            'success': False,
            'message': str(e)
        }, status=500)



@login_required
@role_required(['marketing'])
def create_manual_job(request):
    """Create job manually without AI summary"""
    
    # Get form options
    templates = [
        template for template in TemplateMaster.objects.filter(status='active')
        if not getattr(template, 'is_deleted', False)
    ]
    project_groups = [
        group for group in ProjectGroupMaster.objects.all()
        if not getattr(group, 'is_deleted', False)
    ]
    referencing_styles = [
        ref for ref in ReferencingMaster.objects.all()
        if not getattr(ref, 'is_deleted', False)
    ]
    writing_styles = [
        writing for writing in AcademicWritingMaster.objects.all()
        if not getattr(writing, 'is_deleted', False)
    ]
    price_entries = [
        entry for entry in PriceMaster.objects.all()
        if not getattr(entry, 'is_deleted', False)
    ]
    
    # Build price map
    price_map = {}
    for entry in price_entries:
        category_key = str(entry.category).upper()
        level_key = str(entry.level).lower()
        level_label = entry.get_level_display() if hasattr(entry, 'get_level_display') else level_key.title()
        price_map.setdefault(category_key, {})[level_key] = {
            'price': _to_float(entry.price_per_word),
            'label': level_label,
        }
    
    # Get customers
    customers_qs = Customer.objects.filter(created_by=request.user).order_by('customer_name')
    customers = [c for c in customers_qs if getattr(c, 'is_active', False)]
    
    context = {
        'user': request.user,
        'templates': templates,
        'project_groups': project_groups,
        'referencing_styles': referencing_styles,
        'writing_styles': writing_styles,
        'category_choices': Job.CATEGORY_CHOICES,
        'referencing_choices': Job.REFERENCING_STYLE_CHOICES,
        'writing_choices': Job.WRITING_STYLE_CHOICES,
        'level_choices': Job.LEVEL_CHOICES,
        'price_map': json.dumps(price_map),
        'customers': customers,
    }
    
    return render(request, 'marketing/create_manual_job.html', context)


@login_required
@role_required(['marketing'])
@require_http_methods(["POST"])
def submit_manual_job(request):
    """Submit manual job form"""
    try:
        # Extract all form data
        job_id = request.POST.get('job_id', '').strip()
        instruction = request.POST.get('instruction', '').strip()
        topic = request.POST.get('topic', '').strip()
        customer_id = request.POST.get('customer_id', '').strip()
        word_count = request.POST.get('word_count', '').strip()
        referencing_style = request.POST.get('referencing_style', '').strip()
        writing_style = request.POST.get('writing_style', '').strip()
        category = request.POST.get('category', '').strip()
        level = request.POST.get('level', '').strip()
        software = request.POST.get('software', '').strip()
        template_id = request.POST.get('template', '').strip()
        project_group_id = request.POST.get('project_group', '').strip()
        expected_deadline = request.POST.get('expected_deadline', '').strip()
        strict_deadline = request.POST.get('strict_deadline', '').strip()
        amount = request.POST.get('amount', '').strip()
        files = request.FILES.getlist('attachments')
        
        # Validation
        errors = []
        
        if not job_id:
            errors.append('Job ID is required.')
        else:
            # Check uniqueness
            if Job.objects.filter(job_id=job_id).exists():
                errors.append('Job ID already exists.')
        
        if not instruction or len(instruction) < 50:
            errors.append('Instruction must be at least 50 characters.')
        
        if not topic:
            errors.append('Topic is required.')
        
        if not customer_id:
            errors.append('Customer selection is required.')
        
        if not word_count or not word_count.isdigit():
            errors.append('Valid word count is required.')
        else:
            word_count = int(word_count)
            if word_count <= 0:
                errors.append('Word count must be greater than 0.')
        
        if not category:
            errors.append('Category is required.')
        
        if not level:
            errors.append('Level is required.')
        
        if not template_id:
            errors.append('Template selection is required.')
        
        if not project_group_id:
            errors.append('Project group selection is required.')
        
        if not expected_deadline:
            errors.append('Expected deadline is required.')
        
        if not strict_deadline:
            errors.append('Strict deadline is required.')
        
        if not amount:
            errors.append('Amount is required.')
        else:
            try:
                amount = float(amount)
                if amount <= 0:
                    errors.append('Amount must be greater than 0.')
            except ValueError:
                errors.append('Invalid amount format.')
        
        # Validate files
        if files:
            if len(files) > 10:
                errors.append('Maximum 10 files allowed.')
            for file in files:
                is_valid, msg = validate_file(file)
                if not is_valid:
                    errors.append(msg)
        
        # Validate deadlines
        if expected_deadline and strict_deadline:
            try:
                expected_dt = timezone.datetime.fromisoformat(expected_deadline)
                strict_dt = timezone.datetime.fromisoformat(strict_deadline)
                
                if timezone.is_naive(expected_dt):
                    expected_dt = timezone.make_aware(expected_dt)
                if timezone.is_naive(strict_dt):
                    strict_dt = timezone.make_aware(strict_dt)
                
                min_strict = timezone.now() + timedelta(hours=24)
                if strict_dt < min_strict:
                    errors.append('Strict deadline must be at least 24 hours from now.')
                
                if expected_dt >= strict_dt:
                    errors.append('Expected deadline should be before strict deadline.')
                
            except ValueError:
                errors.append('Invalid deadline format.')
        
        if errors:
            return JsonResponse({
                'success': False,
                'errors': errors
            }, status=400)
        
        # Get customer
        customer_obj = get_object_or_404(
            Customer,
            customer_id=customer_id,
            created_by=request.user
        )
        
        # Get related objects
        template = get_object_or_404(TemplateMaster, id=template_id)
        if getattr(template, 'is_deleted', False):
            return JsonResponse({
                'success': False,
                'errors': ['Selected template is no longer available.']
            }, status=400)

        project_group = get_object_or_404(ProjectGroupMaster, id=project_group_id)
        if getattr(project_group, 'is_deleted', False):
            return JsonResponse({
                'success': False,
                'errors': ['Selected project group is no longer available.']
            }, status=400)
        
        # Calculate system expected amount
        normalized_category = category.upper()
        price_level = _to_price_master_level(level)
        price_entry = next(
            (
                item for item in PriceMaster.objects.filter(
                    category=normalized_category,
                    level=price_level
                )
                if not getattr(item, 'is_deleted', False)
            ),
            None
        )
        
        if not price_entry:
            return JsonResponse({
                'success': False,
                'errors': ['Pricing not configured for the selected category and level.']
            }, status=400)
        
        price_per_word = _to_float(price_entry.price_per_word)
        system_expected = price_per_word * float(word_count)
        normalized_level = _normalize_level(level)
        
        with transaction.atomic():
            # Generate system ID
            system_id = Job.generate_system_id()
            masking_generated_at = timezone.now()
            
            # Create job
            job = Job.objects.create(
                system_id=system_id,
                job_id=job_id,
                instruction=instruction,
                topic=topic,
                category=category,
                level=normalized_level,
                word_count=word_count,
                referencing_style=referencing_style or None,
                writing_style=writing_style or None,
                software=software or None,
                template=template,
                project_group=project_group,
                expected_deadline=expected_dt,
                strict_deadline=strict_dt,
                amount=amount,
                system_expected_amount=system_expected,
                customer_name=customer_obj.customer_name,
                customer_id=customer_obj.customer_id,
                created_by=request.user,
                status='unallocated',
                job_name_validated_at=timezone.now(),
                initial_form_submitted_at=timezone.now(),
                final_form_submitted_at=timezone.now(),
                masking_id_generated_at=masking_generated_at,
            )
            
            # Save attachments if any
            for file in files:
                JobAttachment.objects.create(
                    job=job,
                    file=file,
                    original_filename=file.name,
                    file_size=file.size,
                    uploaded_by=request.user
                )
            
            # Generate project prefix and masking ID
            project_prefix = JobTemplate.generate_project_prefix(project_group)
            masking_id = JobTemplate.generate_masking_id(category, project_prefix)
            
            # Create JobTemplate instance
            job_template = JobTemplate.objects.create(
                job=job,
                template=template,
                masking_id=masking_id,
                project_group=project_group,
                project_prefix=project_prefix,
                created_by=request.user,
                masking_id_generated_at=masking_generated_at
            )
            
            # Create default tasks
            task_configs = template.default_tasks
            for task_config in task_configs:
                task_id = JobTask.generate_task_id(
                    project_prefix,
                    task_config['task_code']
                )
                
                JobTask.objects.create(
                    job_template=job_template,
                    task_id=task_id,
                    task_number=task_config['task_number'],
                    task_name=task_config['task_name'],
                    task_code=task_config['task_code'],
                    assignable_roles=task_config['assignable_roles'],
                    status='pending',
                    word_count=word_count if task_config['task_number'] == 1 else None
                )
            
            # Log activities
            JobActionLog.objects.create(
                job=job,
                action='created',
                performed_by=request.user,
                performed_by_type='user',
                details={
                    'method': 'manual',
                    'job_id': job_id,
                    'topic': topic,
                    'category': category,
                    'level': normalized_level,
                }
            )
            
            ActivityLog.objects.create(
                event_key='job.created.manual',
                category='job_management',
                subject_user=request.user,
                performed_by=request.user,
                metadata={
                    'job_system_id': system_id,
                    'job_id': job_id,
                    'masking_id': masking_id,
                    'method': 'manual',
                    'status': 'unallocated',
                }
            )
            
            return JsonResponse({
                'success': True,
                'message': f'Job "{job_id}" created successfully!',
                'system_id': system_id,
                'masking_id': masking_id,
                'redirect': reverse('marketing_dashboard')
            })
        
    except Exception as e:
        logger.exception(f"Error creating manual job: {str(e)}")
        return JsonResponse({
            'success': False,
            'errors': [f'An error occurred: {str(e)}']
        }, status=500)
